from flask import Flask, request, render_template, send_file
from docx import Document
from docx.shared import Pt
import os

app = Flask(__name__)

# Función para reemplazar texto en negrita
def reemplazar_con_negrita(doc, data):
    for p in doc.paragraphs:
        for key, value in data.items():
            marcador = f"{{{key}}}"
            if marcador in p.text:
                nueva_parte = []
                for run in p.runs:
                    if marcador in run.text:
                        partes = run.text.split(marcador)
                        for i, parte in enumerate(partes):
                            nueva_parte.append((parte, run.bold))  # mantener formato
                            if i < len(partes) - 1:
                                nueva_parte.append((value, True))   # nuevo valor en negrita
                        run.text = ""
                    else:
                        nueva_parte.append((run.text, run.bold))
                        run.text = ""
                for texto, bold in nueva_parte:
                    new_run = p.add_run(texto)
                    new_run.bold = bold
                    new_run.font.size = Pt(11)

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        # Captura de datos del formulario
        data = {
            "nombreRRLL": request.form["nombreRRLL"],
            "dni": request.form["dni"],
            "fechaylugarnacimiento": request.form["fechaylugarnacimiento"],
            "domiciliofiscal": request.form["domiciliofiscal"],
            "direccioninstalacion": request.form["direccioninstalacion"],
            "nombrepadres": request.form["nombrepadres"],
            "razonsocial": request.form["razonsocial"],
            "ruc": request.form["ruc"],
            "cargo": request.form["cargo"],
            "nombreGerente": request.form["nombreGerente"],
            "numero": request.form["numero"],
            "correo": request.form["correo"],
            "paginasblancas": request.form["paginasblancas"]
        }
        print(data)
        speech = request.form["speech"]

        ruta_speech = f"./speeches/{speech}.docx"
        if not os.path.exists(ruta_speech):
            return f"❌ Plantilla para el paquete '{speech}' no encontrada.", 404

        # Cargar documento y reemplazar variables
        doc = Document(ruta_speech)
        reemplazar_con_negrita(doc, data)

        # Guardar documento generado
        razon = data["razonsocial"]
        ruc = data["ruc"]
        nombre_archivo = f"{razon}_{ruc}_{speech}.docx"
        output_path = f"./contratos_generados/{nombre_archivo}"
        doc.save(output_path)

        # Descargar el archivo
        return send_file(output_path, as_attachment=True, download_name=nombre_archivo)

    return render_template("formulario.html")

if __name__ == "__main__":
    os.makedirs("contratos_generados", exist_ok=True)
    app.run(debug=True)
